from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.schemas import GenerationResponse, ProgramAnalysis, Question
from app.services.discipline_catalog import find_best_profile, load_catalog_profiles


EXPORT_DIR = Path(__file__).resolve().parents[1] / "storage" / "exports"
EXPORT_DIR.mkdir(parents=True, exist_ok=True)

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(14)
BLACK = RGBColor(0, 0, 0)

TYPE_LABELS = {
    "open": "теоретический вопрос",
    "practice": "практическое задание",
    "test": "тестовое задание",
}

DIFFICULTY_LABELS = {
    "easy": "базовый уровень",
    "medium": "средний уровень",
    "hard": "повышенный уровень",
}


def export_generation_to_docx(generation: GenerationResponse, program: ProgramAnalysis | None = None) -> Path:
    """Export a control work as a clean printable DOCX.

    Word heading styles are preserved for the navigation pane. All visible text is
    forced to Times New Roman 14 pt, 1.5 spacing and black color.
    """
    document = Document()
    _configure_document(document)

    discipline_name = _detect_discipline_name(program)
    variants_count = len(generation.variants)
    questions_total = sum(len(variant.questions) for variant in generation.variants)

    _add_title_page(document, discipline_name, variants_count, questions_total)
    _add_student_variants(document, generation)
    _add_answer_appendix(document, generation)
    _force_black_times_new_roman(document)

    output_path = EXPORT_DIR / f"control_work_{generation.session_id}.docx"
    document.save(output_path)
    return output_path


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3", "List Bullet"]:
        style = document.styles[style_name]
        style.font.name = FONT_NAME
        style.font.size = FONT_SIZE
        style.font.color.rgb = BLACK
        _set_rfonts(style.element, FONT_NAME)
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph_format.line_spacing = 1.5
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)

    document.styles["Heading 1"].font.bold = True
    document.styles["Heading 2"].font.bold = True
    document.styles["Heading 3"].font.bold = True


def _add_title_page(document: Document, discipline_name: str, variants_count: int, questions_total: int) -> None:
    title = document.add_heading("Контрольная работа", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    discipline = document.add_paragraph()
    discipline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_bold_run(discipline, f"по дисциплине «{discipline_name}»")

    _add_blank(document)
    _add_label_value(document, "Форма контроля", "контрольная работа")
    _add_label_value(document, "Количество вариантов", str(variants_count))
    _add_label_value(document, "Количество заданий", str(questions_total))
    _add_label_value(document, "Дата формирования", datetime.now().strftime("%d.%m.%Y"))

    _add_blank(document)
    instruction_heading = document.add_heading("Инструкция по выполнению", level=2)
    _format_paragraph(instruction_heading)
    instructions = [
        "Внимательно прочитайте формулировку каждого задания.",
        "Ответы должны быть развернутыми, логически связанными с темой дисциплины и сопровождаться примерами применения.",
        "При выполнении практических заданий необходимо указать цель, входные данные, последовательность действий и ожидаемый результат.",
        "Оформляйте ответы аккуратно, соблюдая терминологию дисциплины.",
    ]
    for item in instructions:
        paragraph = document.add_paragraph(item, style="List Bullet")
        _format_paragraph(paragraph)

    document.add_page_break()


def _add_student_variants(document: Document, generation: GenerationResponse) -> None:
    section_heading = document.add_heading("Задания для выполнения", level=1)
    section_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for variant_index, variant in enumerate(generation.variants):
        if variant_index:
            document.add_page_break()
        heading = document.add_heading(f"Вариант {variant.variant_number}", level=2)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for index, question in enumerate(variant.questions, start=1):
            _add_question_block(document, index, question, include_answer=False)


def _add_answer_appendix(document: Document, generation: GenerationResponse) -> None:
    document.add_page_break()
    heading = document.add_heading("Приложение А. Эталонные ответы и критерии оценивания", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for variant in generation.variants:
        variant_heading = document.add_heading(f"Вариант {variant.variant_number}", level=2)
        variant_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for index, question in enumerate(variant.questions, start=1):
            _add_question_block(document, index, question, include_answer=True)


def _add_question_block(document: Document, index: int, question: Question, include_answer: bool) -> None:
    heading = document.add_heading(f"Задание {index}", level=3)
    _format_paragraph(heading)

    task = document.add_paragraph()
    _add_bold_run(task, "Формулировка: ")
    task.add_run(_normalize_question_text(question))
    _format_paragraph(task)

    meta = document.add_paragraph()
    meta.add_run("Тема: ")
    _add_bold_run(meta, _clean_text(question.topic))
    meta.add_run(". Тип задания: ")
    _add_bold_run(meta, TYPE_LABELS.get(question.type, _clean_text(question.type)))
    meta.add_run(". Уровень сложности: ")
    _add_bold_run(meta, DIFFICULTY_LABELS.get(question.difficulty, _clean_text(question.difficulty)))
    meta.add_run(".")
    _format_paragraph(meta)

    if include_answer:
        answer = _clean_text(question.answer)
        if answer:
            answer_paragraph = document.add_paragraph()
            _add_bold_run(answer_paragraph, "Эталонный ответ: ")
            answer_paragraph.add_run(answer)
            _format_paragraph(answer_paragraph)

        criteria = [_clean_text(item) for item in question.criteria if _clean_text(item)]
        if criteria:
            criteria_heading = document.add_paragraph()
            _add_bold_run(criteria_heading, "Критерии оценивания:")
            _format_paragraph(criteria_heading)
            for criterion in criteria:
                criterion_paragraph = document.add_paragraph(criterion, style="List Bullet")
                _format_paragraph(criterion_paragraph)

    _add_blank(document)


def _add_label_value(document: Document, label: str, value: str) -> None:
    paragraph = document.add_paragraph()
    _add_bold_run(paragraph, f"{label}: ")
    paragraph.add_run(value)
    _format_paragraph(paragraph)


def _add_blank(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)


def _add_bold_run(paragraph, text: str):
    run = paragraph.add_run(text)
    run.bold = True
    _format_run(run)
    return run


def _format_paragraph(paragraph) -> None:
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        _format_run(run)


def _format_run(run) -> None:
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.font.color.rgb = BLACK
    _set_rfonts(run._element, FONT_NAME)


def _set_rfonts(element, font_name: str) -> None:
    r_pr = element.find(qn("w:rPr"))
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        element.insert(0, r_pr)
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:cs"), font_name)
    r_fonts.set(qn("w:eastAsia"), font_name)


def _force_black_times_new_roman(document: Document) -> None:
    for paragraph in document.paragraphs:
        _format_paragraph(paragraph)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _format_paragraph(paragraph)


def _normalize_question_text(question: Question) -> str:
    text = _clean_text(question.text)
    topic = _clean_text(question.topic)
    if _looks_malformed(text):
        if question.type == "practice":
            return f"Выполните практическое задание по теме «{topic}». Опишите цель применения, входные данные, последовательность действий и ожидаемый результат."
        if question.type == "test":
            return f"Выберите правильный ответ по теме «{topic}» и кратко обоснуйте выбор."
        return f"Раскройте содержание темы «{topic}», укажите ее практическое назначение и приведите пример применения."
    return text


def _looks_malformed(text: str) -> bool:
    if not text:
        return True
    if text[:1] in {":", ",", ";", "."}:
        return True
    if "\ufffe" in text or "\ufffd" in text:
        return True
    if re.search(r"^[A-ZА-Я]{1,4}\.?\s*[:.]\s+", text):
        return True
    return False


def _clean_text(value: str) -> str:
    text = str(value or "")
    text = text.replace("\ufffe", "").replace("\ufffd", "")
    text = text.replace("\u00a0", " ")
    text = re.sub(r"\s+", " ", text).strip()
    text = re.sub(r"\s+([,.;:!?])", r"\1", text)
    text = re.sub(r"([,.;:!?])(?=[^\s»)])", r"\1 ", text)
    text = text.replace(" - ", "-")
    text = text.replace("веб прилож", "веб-прилож")
    text = text.strip(" ,;:")
    return text


def _detect_discipline_name(program: ProgramAnalysis | None) -> str:
    if program is None:
        return "не указана"

    from_text = _extract_discipline_from_text(program.text_preview)
    if from_text:
        return from_text

    match = find_best_profile(program.filename, program.text_preview, program.topics, load_catalog_profiles())
    if match:
        profile, _score = match
        discipline_name = _clean_text(profile.get("discipline_name", ""))
        if discipline_name:
            return discipline_name

    filename = Path(program.filename).stem.replace("_", " ").strip()
    return filename or "не указана"


def _extract_discipline_from_text(text: str) -> str:
    value = _clean_text(text)
    patterns = [
        r"(?:дисциплина|дисциплины)\s*[«\"]([^»\"]{4,120})[»\"]",
        r"(?:наименование\s+дисциплины|название\s+дисциплины)\s*[:\-]\s*([^.;\n]{4,120})",
        r"по\s+дисциплине\s*[«\"]([^»\"]{4,120})[»\"]",
    ]
    for pattern in patterns:
        match = re.search(pattern, value, flags=re.IGNORECASE)
        if match:
            return _clean_text(match.group(1))
    return ""
