from __future__ import annotations

import re
from collections import defaultdict
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from app.assessment_item_validation import AssessmentItemsValidation
from app.schemas import AssessmentFundResponse, AssessmentFundSection, AssessmentItemRead

EXPORT_DIR = Path(__file__).resolve().parents[1] / "storage" / "exports"
EXPORT_DIR.mkdir(parents=True, exist_ok=True)

SPECIAL_SECTION_TYPES = {"competency_matrix", "grading_rubric"}
CURRENT_CONTROL_TYPES = {"oral", "practice", "laboratory", "test_bank", "report_topics", "control_work"}
INTERMEDIATE_CONTROL_TYPES = {"exam_questions", "exam_practice", "credit"}


def export_assessment_fund_to_docx(
    fund: AssessmentFundResponse,
    items: list[AssessmentItemRead],
    validation: AssessmentItemsValidation,
) -> Path:
    document = Document()
    _configure_document(document)
    _add_title_page(document, fund)
    _add_competency_matrix(document, fund)
    _add_assessment_procedures(document, fund, items)
    _add_grading_rubric(document)
    _add_answers_appendix(document, fund, items)
    _add_validation_appendix(document, validation)

    filename = f"fos_{_safe_filename(fund.discipline_name)}_{fund.fund_id[:8]}.docx"
    output_path = EXPORT_DIR / filename
    document.save(output_path)
    return output_path


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(14)
        style.font.bold = True
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(6)


def _add_title_page(document: Document, fund: AssessmentFundResponse) -> None:
    _add_centered(document, "Федеральное государственное бюджетное образовательное учреждение", size=14)
    _add_centered(document, "высшего образования", size=14)
    _add_centered(document, "«Саратовский государственный технический университет имени Гагарина Ю.А.»", size=14)
    _add_centered(document, "Кафедра", size=14, space_before=8)
    _add_centered(document, "Оценочные материалы по дисциплине", bold=True, size=16, space_before=80)
    _add_centered(document, f"«{_clean_doc_text(fund.discipline_name)}»", bold=True, size=16, space_before=12)
    _add_centered(document, "направления подготовки", size=14, space_before=12)
    _add_centered(document, "09.03.02 Информационные системы и технологии", size=14)
    _add_centered(document, "Профиль", size=14, space_before=12)
    _add_centered(document, "«Информационные системы и технологии»", size=14)
    _add_centered(document, f"Саратов {datetime.now().year}", size=14, space_before=180)
    document.add_page_break()


def _add_competency_matrix(document: Document, fund: AssessmentFundResponse) -> None:
    _add_heading(document, "1. Перечень компетенций и уровни их сформированности по дисциплине", level=1)
    _add_body(
        document,
        f"В процессе освоения образовательной программы у обучающегося в ходе изучения дисциплины «{_clean_doc_text(fund.discipline_name)}» должны сформироваться компетенции: {_competency_codes(fund)}.",
    )
    _add_body(document, "Критерии определения сформированности компетенций на различных уровнях их формирования приведены ниже.")

    if not fund.competencies:
        _add_body(document, "Компетенции не распознаны автоматически. Раздел требует ручного заполнения преподавателем или методистом.", italic=True)
        return

    for competency in fund.competencies:
        _add_competency_block(document, competency)


def _add_competency_block(document: Document, competency) -> None:
    summary = document.add_table(rows=1, cols=2)
    summary.style = "Table Grid"
    summary.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_cell_text(summary.rows[0].cells[0], "Индекс компетенции", bold=True)
    _set_cell_text(summary.rows[0].cells[1], "Содержание компетенции", bold=True)
    row = summary.add_row().cells
    _set_cell_text(row[0], _clean_doc_text(competency.code))
    _set_cell_text(row[1], _clean_doc_text(competency.description or f"Компетенция {competency.code}, формируемая в рамках освоения дисциплины."))

    indicator_table = document.add_table(rows=1, cols=3)
    indicator_table.style = "Table Grid"
    indicator_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, value in enumerate(["Код и наименование индикатора достижения компетенции", "Виды занятий для формирования компетенции", "Оценочные средства"]):
        _set_cell_text(indicator_table.rows[0].cells[index], value, bold=True)
    indicators = competency.indicators or [f"Индикатор достижения {competency.code} требует уточнения."]
    for indicator in indicators[:3]:
        cells = indicator_table.add_row().cells
        _set_cell_text(cells[0], _clean_doc_text(indicator))
        _set_cell_text(cells[1], "лекции, практические занятия, самостоятельная работа")
        _set_cell_text(cells[2], "устный опрос, практические задания, промежуточная аттестация, диагностическая работа")

    _add_body(document, "Уровни освоения компетенции", bold=True, first_line=False)
    levels = document.add_table(rows=1, cols=2)
    levels.style = "Table Grid"
    levels.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_cell_text(levels.rows[0].cells[0], "Уровень освоения компетенции", bold=True)
    _set_cell_text(levels.rows[0].cells[1], "Критерии оценивания", bold=True)
    for level, criterion in _competency_level_rows(competency):
        cells = levels.add_row().cells
        _set_cell_text(cells[0], level)
        _set_cell_text(cells[1], criterion)
    document.add_paragraph()


def _competency_level_rows(competency) -> list[tuple[str, str]]:
    base = _clean_doc_text(competency.description or f"содержание компетенции {competency.code}").rstrip(".;")
    return [
        ("Продвинутый (отлично)", f"Обучающийся полно и аргументированно раскрывает {base}, самостоятельно применяет знания при решении учебных и профессионально ориентированных задач."),
        ("Повышенный (хорошо)", f"Обучающийся раскрывает основные положения компетенции, применяет материал при решении типовых задач, допускает отдельные неточности."),
        ("Пороговый (удовлетворительно)", f"Обучающийся демонстрирует минимально допустимое освоение компетенции {competency.code}, выполняет типовые задания при наличии методической поддержки."),
    ]


def _add_assessment_procedures(
    document: Document,
    fund: AssessmentFundResponse,
    items: list[AssessmentItemRead],
) -> None:
    _add_heading(document, "2. Методические, оценочные материалы и средства, определяющие процедуры оценивания сформированности компетенций", level=1)
    _add_body(document, "Оценивание сформированности компетенций проводится в процессе текущего контроля и промежуточной аттестации. Оценочные материалы сгруппированы по видам контроля и темам дисциплины.")

    items_by_section: dict[str, list[AssessmentItemRead]] = defaultdict(list)
    for item in items:
        items_by_section[item.section_code].append(item)

    _add_heading(document, "2.1 Оценочные средства для текущего контроля", level=2)
    _add_sections_by_type(document, fund, items_by_section, CURRENT_CONTROL_TYPES)

    _add_heading(document, "2.2 Оценочные средства для промежуточного контроля", level=2)
    _add_sections_by_type(document, fund, items_by_section, INTERMEDIATE_CONTROL_TYPES)

    _add_heading(document, "2.3 Итоговая диагностическая работа по дисциплине", level=2)
    diagnostic_sections = [section for section in fund.sections if section.enabled and section.assessment_type == "diagnostic"]
    if diagnostic_sections:
        for section in diagnostic_sections:
            section_items = items_by_section.get(section.code, [])
            if section_items:
                _add_diagnostic_blocks(document, section_items)
            else:
                _add_body(document, "Диагностические задания пока не сформированы.", italic=True)
    else:
        _add_body(document, "Итоговая диагностическая работа не включена в структуру ФОС.", italic=True)


def _add_sections_by_type(
    document: Document,
    fund: AssessmentFundResponse,
    items_by_section: dict[str, list[AssessmentItemRead]],
    allowed_types: set[str],
) -> None:
    sections = [section for section in fund.sections if section.enabled and section.assessment_type in allowed_types]
    if not sections:
        _add_body(document, "Разделы данного вида контроля не сформированы автоматически.", italic=True)
        return

    for section in sections:
        section_items = items_by_section.get(section.code, [])
        _add_heading(document, section.title, level=3)
        if not section_items:
            _add_body(document, "Задания для данного раздела пока не сформированы.", italic=True)
            continue
        _add_grouped_items(document, section_items)


def _add_grouped_items(document: Document, items: list[AssessmentItemRead]) -> None:
    grouped: dict[str, list[AssessmentItemRead]] = defaultdict(list)
    for item in items:
        grouped[_clean_doc_text(item.topic or "Тема не указана")].append(item)

    for topic_index, (topic, topic_items) in enumerate(grouped.items(), start=1):
        _add_body(document, f"Тема {topic_index}. {topic}", bold=True, first_line=False)
        for item_index, item in enumerate(topic_items, start=1):
            _add_body(document, f"{item_index}. {_clean_doc_text(item.text)}", first_line=False)


def _add_diagnostic_blocks(document: Document, items: list[AssessmentItemRead]) -> None:
    _add_body(document, "Итоговая диагностическая работа выполняется в форме набора заданий. Для каждого задания указывается проверяемая компетенция, содержание вопроса, варианты ответа и эталонный ответ.")
    for index, item in enumerate(items, start=1):
        _add_body(document, f"Задание {index}.", bold=True, first_line=False)
        _add_body(document, f"Содержание вопроса: {_clean_doc_text(item.text)}", first_line=False)
        if item.criteria:
            _add_body(document, "Варианты ответа:", bold=True, first_line=False)
            for option_index, option in enumerate(item.criteria[:4], start=1):
                _add_body(document, f"{option_index}) {_clean_doc_text(option)}", first_line=False)
        _add_body(document, f"Эталонный ответ: {_clean_doc_text(item.answer or 'Требует уточнения преподавателем.')}", first_line=False)
        _add_body(document, f"Компетенция: {_clean_doc_text(item.competency_code or 'не указана')}; индикатор: {_clean_doc_text(item.indicator or 'не указан')}", first_line=False)


def _criteria_as_options(item: AssessmentItemRead) -> str:
    if not item.criteria:
        return "—"
    return "\n".join(f"{index}. {_clean_doc_text(criterion)}" for index, criterion in enumerate(item.criteria[:5], start=1))


def _add_grading_rubric(document: Document) -> None:
    _add_heading(document, "3. Критерии выставления оценок при проведении текущего контроля и промежуточной аттестации", level=1)
    _add_body(document, "Оценивание результатов обучения проводится путем контроля сформированности элементов компетенций. Оценка выставляется с учетом полноты ответа, корректности выполнения практической части, самостоятельности решения и качества обоснования.")
    _add_body(document, "Оценки «неудовлетворительно» также ставятся при обнаружении списывания, плагиата, фальсификации данных и результатов работы.")

    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Шкала оценки", "Оценка", "Критерий выставления оценки"]
    for index, value in enumerate(headers):
        _set_cell_text(table.rows[0].cells[index], value, bold=True)
    rows = [
        ("Пятибалльная шкала", "Отлично", "Обучающийся ответил на теоретические вопросы, показал глубокие знания учебного материала, полностью выполнил практические задания и продемонстрировал уверенное владение навыками применения знаний при решении задач."),
        ("Пятибалльная шкала", "Хорошо", "Обучающийся ответил на теоретические вопросы, показал знания в рамках учебного материала, выполнил практические задания с отдельными неточностями и продемонстрировал хорошие умения решения задач."),
        ("Пятибалльная шкала", "Удовлетворительно", "Обучающийся показал базовые знания, частично выполнил практические задания и продемонстрировал минимально допустимый уровень применения знаний и умений."),
        ("Пятибалльная шкала", "Неудовлетворительно", "Обучающийся продемонстрировал недостаточный уровень знаний и умений, допустил существенные ошибки или не приступал к выполнению задания."),
    ]
    for scale, mark, criterion in rows:
        cells = table.add_row().cells
        _set_cell_text(cells[0], scale)
        _set_cell_text(cells[1], mark)
        _set_cell_text(cells[2], criterion)
    document.add_paragraph()


def _add_answers_appendix(
    document: Document,
    fund: AssessmentFundResponse,
    items: list[AssessmentItemRead],
) -> None:
    document.add_page_break()
    _add_heading(document, "Приложение А. Эталонные ответы и критерии оценивания", level=1)
    items_by_section: dict[str, list[AssessmentItemRead]] = defaultdict(list)
    for item in items:
        items_by_section[item.section_code].append(item)

    for section in _enabled_content_sections(fund.sections):
        section_items = items_by_section.get(section.code, [])
        if not section_items:
            continue
        _add_heading(document, section.title, level=2)
        for index, item in enumerate(section_items, start=1):
            _add_body(document, f"{index}. {_clean_doc_text(item.text)}", first_line=False)
            _add_body(document, f"Эталонный ответ: {_clean_doc_text(item.answer or 'Требует уточнения преподавателем.')}", first_line=False)
            criteria = [_clean_doc_text(value) for value in (item.criteria or []) if _clean_doc_text(value)]
            if criteria:
                _add_body(document, "Критерии оценивания:", first_line=False, bold=True)
                for criterion in criteria[:4]:
                    paragraph = document.add_paragraph(criterion, style="List Bullet")
                    paragraph.paragraph_format.first_line_indent = Cm(0)
                    paragraph.paragraph_format.left_indent = Cm(0.75)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _add_validation_appendix(document: Document, validation: AssessmentItemsValidation) -> None:
    document.add_page_break()
    _add_heading(document, "Приложение Б. Отчет автоматизированной проверки банка заданий", level=1)
    _add_body(document, "Отчет сформирован автоматически и предназначен для предварительной проверки полноты банка заданий перед экспертной оценкой преподавателем и методистом.")

    metrics = [
        ("Всего заданий", str(validation.total_items)),
        ("Покрытие тем", f"{validation.topics_coverage_score}%"),
        ("Покрытие компетенций", f"{validation.competencies_coverage_score}%"),
        ("Готовность эталонных ответов", f"{validation.answers_readiness_score}%"),
        ("Готовность критериев оценивания", f"{validation.criteria_readiness_score}%"),
        ("Доля потенциальных дублей", f"{validation.duplicate_rate}%"),
    ]
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_cell_text(table.rows[0].cells[0], "Показатель", bold=True)
    _set_cell_text(table.rows[0].cells[1], "Значение", bold=True)
    for label, value in metrics:
        row = table.add_row().cells
        _set_cell_text(row[0], label)
        _set_cell_text(row[1], value)

    if validation.warnings:
        _add_heading(document, "Предупреждения", level=2)
        for warning in validation.warnings:
            paragraph = document.add_paragraph(_clean_doc_text(warning), style="List Bullet")
            paragraph.paragraph_format.first_line_indent = Cm(0)

    _add_heading(document, "Матрица покрытия тем", level=2)
    coverage = document.add_table(rows=1, cols=4)
    coverage.style = "Table Grid"
    coverage.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Тема", "Количество заданий", "Разделы", "Компетенции"]
    for index, value in enumerate(headers):
        _set_cell_text(coverage.rows[0].cells[index], value, bold=True)
    for row_data in validation.coverage_rows:
        row = coverage.add_row().cells
        sections = "; ".join(f"{key}: {value}" for key, value in row_data.section_counts.items()) or "—"
        _set_cell_text(row[0], _clean_doc_text(row_data.topic))
        _set_cell_text(row[1], str(row_data.total_items))
        _set_cell_text(row[2], sections)
        _set_cell_text(row[3], ", ".join(row_data.competencies) or "—")

    if validation.duplicate_groups:
        _add_heading(document, "Потенциальные дубли", level=2)
        for index, group in enumerate(validation.duplicate_groups, start=1):
            _add_body(document, f"Группа {index}; сходство {round(group.similarity * 100)}%; связанных заданий: {len(group.item_ids)}. Пример формулировки: {_clean_doc_text(group.sample_text)}", first_line=False)


def _enabled_content_sections(sections: list[AssessmentFundSection]) -> list[AssessmentFundSection]:
    return [
        section
        for section in sections
        if section.enabled and section.assessment_type not in SPECIAL_SECTION_TYPES and section.assessment_type != "diagnostic"
    ]


def _competency_codes(fund: AssessmentFundResponse) -> str:
    codes = [competency.code for competency in fund.competencies]
    return ", ".join(codes) if codes else "не определены автоматически"


def _add_heading(document: Document, text: str, level: int, centered: bool = False) -> None:
    paragraph = document.add_heading(_clean_doc_text(text), level=level)
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT


def _add_body(document: Document, text: str, *, italic: bool = False, bold: bool = False, first_line: bool = True) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(1.25) if first_line else Cm(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run(_clean_doc_text(text))
    run.italic = italic
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14)


def _add_centered(document: Document, text: str, *, bold: bool = False, size: int = 14, space_before: int = 0) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(space_before)
    run = paragraph.add_run(_clean_doc_text(text))
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)


def _set_cell_text(cell, value: str, *, bold: bool = False) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(_clean_doc_text(value))
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(10)
    _set_cell_margins(cell, top=60, start=80, bottom=60, end=80)


def _set_cell_margins(cell, **kwargs: int) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin in ("top", "start", "bottom", "end"):
        if margin in kwargs:
            node = tc_mar.find(qn(f"w:{margin}"))
            if node is None:
                node = OxmlElement(f"w:{margin}")
                tc_mar.append(node)
            node.set(qn("w:w"), str(kwargs[margin]))
            node.set(qn("w:type"), "dxa")


def _clean_doc_text(value: str | None) -> str:
    value = (value or "").replace("\ufffe", "-").replace("\u00ad", "")
    value = value.replace("#default#", "")
    value = re.sub(r"\s+", " ", value).strip(" .;:-—\t\n\r")
    return value


def _safe_filename(value: str) -> str:
    value = re.sub(r"[^0-9A-Za-zА-Яа-яЁё_-]+", "_", value.strip())
    value = re.sub(r"_+", "_", value).strip("_")
    return value[:80] or "discipline"
