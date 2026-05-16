from pathlib import Path

import fitz
from docx import Document


SUPPORTED_EXTENSIONS = {".docx", ".txt", ".pdf"}


class UnsupportedDocumentFormat(ValueError):
    pass


def extract_text(file_path: str | Path) -> str:
    path = Path(file_path)
    extension = path.suffix.lower()

    if extension not in SUPPORTED_EXTENSIONS:
        raise UnsupportedDocumentFormat(
            f"Unsupported file format: {extension}. Supported: DOCX, PDF, TXT."
        )

    if extension == ".docx":
        return _extract_docx(path)

    if extension == ".pdf":
        return _extract_pdf(path)

    return path.read_text(encoding="utf-8", errors="ignore")


def _extract_docx(path: Path) -> str:
    document = Document(str(path))
    parts: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _extract_pdf(path: Path) -> str:
    parts: list[str] = []
    with fitz.open(str(path)) as pdf:
        for page in pdf:
            text = page.get_text("text").strip()
            if text:
                parts.append(text)
    return "\n".join(parts)
